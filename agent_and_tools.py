import os
import shutil
import pyperclip
import pyautogui
from datetime import datetime
from langchain_core.tools import tool
from langchain_community.tools import TavilySearchResults
import platform
import socket
import base64
import getpass
from io import BytesIO
from langchain_google_genai import ChatGoogleGenerativeAI
from dotenv import load_dotenv
from langchain_core.messages import HumanMessage, SystemMessage
import subprocess
import glob
import zipfile
from PIL import Image
import mss
import pygetwindow as gw
import pypdf
import docx
import openpyxl
import requests
from bs4 import BeautifulSoup
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.prompts import PromptTemplate
from docx import Document
from fpdf import FPDF
import pandas as pd


load_dotenv()

# Create llm
llm = ChatGoogleGenerativeAI(
    model = "gemini-2.5-flash-preview-05-20",
    temperature = 0
)

@tool
def get_username() -> str:
    """
    Retrieves the current Windows username of the person 
    running this script. This is typically the name of the 
    user's home directory (e.g., 'C:\\Users\\hp') and can be 
    used for personalization, file path construction, or 
    system identification in LangGraph applications.
    """
    return getpass.getuser()

@tool
def Save_to_clipboard(clipboard_content: str):
    """This is a function to save the content to the user clipboard.
    
    Args:
        clipboard_content: The content to be saved.
    """
    pyperclip.copy(clipboard_content)

    return "Content has been copied to the clipboard."

@tool
def get_clipboard_content() -> str:
    """
    Retrieves the current content from the system clipboard.
    
    Returns:
        The text content of the clipboard.
    """
    try:
        return pyperclip.paste()
    except Exception as e:
        return f"Error reading from clipboard: {str(e)}"

@tool
def list_directory_tree(path: str, depth: int = 0) -> str:
    """
    Returns the folder and file tree of a given directory up to a specified depth.

    Args:
        path: The root directory path to explore.
        depth: How many levels deep to explore (default: 0).
    """
    tree_lines = []

    def walk(current_path, current_depth):
        if current_depth > depth:
            return
        try:
            items = os.listdir(current_path)
        except (PermissionError, FileNotFoundError):
            return
        for item in items:
            item_path = os.path.join(current_path, item)
            tree_lines.append("  " * current_depth + f"- {item}")
            if os.path.isdir(item_path):
                walk(item_path, current_depth + 1)

    if not os.path.exists(path):
        return f"Path does not exist: {path}"
    if not os.path.isdir(path):
        return f"Path is not a directory: {path}"

    walk(path, 0)
    return "\n".join(tree_lines)

@tool
def find_files(start_dir: str, pattern: str) -> str:
    """
    Finds files matching a specific pattern within a directory and its subdirectories.

    Args:
        start_dir: The directory to start the search from.
        pattern: The search pattern to match (e.g., '*.txt', 'report.*', 'image_?.png').
                 Uses standard glob patterns.

    Returns:
        A string listing the matching file paths, or a message if none are found.
    """
    try:
        # Use recursive glob to find all matching files
        results = glob.glob(os.path.join(start_dir, f"**/{pattern}"), recursive=True)
        if not results:
            return f"No files found matching '{pattern}' in '{start_dir}'."
        
        return "Found files:\n" + "\n".join(results)
    except Exception as e:
        return f"An error occurred while searching for files: {str(e)}"

@tool
def Create(path: str, name: str) -> str:
    """
    Create a new file or folder with the specified name at the given path.
    If the name contains a dot (e.g., 'file.txt'), it is treated as a file.
    Otherwise, it is treated as a folder.

    Args:
        path: The directory in which to create the file or folder.
        name: The name of the file (with extension) or folder.
    """
    try:
        full_path = os.path.join(path, name)

        if "." in name:
            # It's a file
            with open(full_path, 'w') as f:
                pass
            return f"File created: {full_path}"
        else:
            # It's a folder
            os.makedirs(full_path, exist_ok=True)
            return f"Folder created: {full_path}"
    except Exception as e:
        return f"Error: {str(e)}"

@tool
def Delete(path: str, name: str) -> str:
    """
    Delete a file or folder by specifying its path and name.

    Args:
        path: The directory where the file or folder is located.
        name: The name of the file or folder to delete.
    """
    try:
        target = os.path.join(path, name)
        if os.path.isfile(target):
            os.remove(target)
            return f"File deleted: {target}"
        elif os.path.isdir(target):
            shutil.rmtree(target)
            return f"Folder deleted: {target}"
        else:
            return "Item not found at specified path."
    except Exception as e:
        return f"Error: {str(e)}"
    
@tool
def Move(source: str, destination: str, name: str) -> str:
    """
    Move a file or folder by specifying its current name and source/destination paths.

    Args:
        source: The directory where the file or folder currently exists.
        destination: The directory where the file or folder should be moved.
        name: The current name of the file or folder to move.
    """
    try:
        source_path = os.path.join(source, name)
        destination_path = os.path.join(destination, name)
        shutil.move(source_path, destination_path)
        return f"Moved to: {destination_path}"
    except Exception as e:
        return f"Error: {str(e)}"

@tool
def Rename(path: str, old_name: str, new_name: str) -> str:
    """
    Rename a file or folder from old_name to new_name within the specified path.

    Args:
        path: The directory containing the file or folder.
        old_name: The current name of the file or folder.
        new_name: The new name to assign.
    """
    try:
        old_path = os.path.join(path, old_name)
        new_path = os.path.join(path, new_name)
        os.rename(old_path, new_path)
        return f"Renamed to: {new_path}"
    except Exception as e:
        return f"Error: {str(e)}"

@tool
def Write(path: str, file_name: str, content: str) -> str:
    """
    Write content to a file in various formats including .txt, .docx, .pdf, and .xlsx.

    This tool enables writing structured or unstructured text to a file in the desired format,
    depending on the extension specified in the file name. It supports:

    - .txt: Saves the content as plain text.
    - .docx: Writes the content to a Microsoft Word document.
    - .pdf: Converts the text content into a PDF file with basic formatting.
    - .xlsx: Stores the content in an Excel spreadsheet, splitting lines into rows and tab-separated fields into columns.

    If the target directory does not exist, it will be created automatically.

    Args:
        path (str): The directory where the file should be created or saved.
        file_name (str): The full name of the file including its extension.
        content (str): The text content to be written into the file.

    Returns:
        str: A success message indicating where the file was saved, or an error message in case of failure.
    """
    try:
        os.makedirs(path, exist_ok=True)
        full_path = os.path.join(path, file_name)
        ext = os.path.splitext(file_name)[-1].lower()

        if ext == '.txt':
            with open(full_path, 'w', encoding='utf-8') as f:
                f.write(content)

        elif ext == '.docx':
            doc = Document()
            for line in content.splitlines():
                doc.add_paragraph(line)
            doc.save(full_path)

        elif ext == '.pdf':
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)
            for line in content.splitlines():
                pdf.multi_cell(0, 10, line)
            pdf.output(full_path)

        elif ext == '.xlsx':
            df = pd.DataFrame([line.split('\t') for line in content.splitlines()])
            df.to_excel(full_path, index=False, header=False)

        else:
            return "Unsupported file type."

        return f"Content written to: {full_path}"

    except Exception as e:
        return f"Error writing file: {str(e)}"

@tool
def Read(path: str, file_name: str) -> str:
    """
    Reads and extracts the text content from various file types, including plain text,
    PDF, Microsoft Word (.docx), and Microsoft Excel (.xlsx, xlsm).

    This tool automatically detects the file type based on its extension and uses the
    appropriate method to extract all readable text.

    Args:
        path: The directory where the file is located.
        file_name: The name of the file (including its extension, e.g., 'report.pdf').

    Returns:
        A string containing all the extracted text from the file.
        Returns an error message if the file is not found or the format is unsupported.
    """

    full_path = os.path.join(path, file_name)
    try:
        _, extension = os.path.splitext(file_name)
        extension = extension.lower()

        content = ""

        # Handle PDF files
        if extension == '.pdf':
            with open(full_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    content += page.extract_text() or ''
            return content

        # Handle Word documents
        elif extension == '.docx':
            doc = docx.Document(full_path)
            for para in doc.paragraphs:
                content += para.text + '\n'
            return content

        # Handle Excel files - reads all cells from all sheets
        elif extension in ['.xlsx', '.xlsm']:
            workbook = openpyxl.load_workbook(full_path)
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value is not None:
                            content += str(cell.value) + ' '
                    content += '\n'
            return content
        
        # Handle plain text files as the default
        elif extension == '.txt':
            with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        
        else:
            return "Unsupported file type."

    except FileNotFoundError:
        return f"Error: The file '{full_path}' was not found."
    except Exception as e:
        return f"Error reading file '{full_path}': {str(e)}"

@tool
def open_file_or_app(path: str) -> str:
    """
    Opens a file with its default application or launches an application.

    This function can be used to:
    - Open a document (e.g., 'C:\\Users\\hp\\Documents\\report.docx')
    - Open an image or video
    - Launch an executable application (e.g., 'C:\\Program Files\\Google\\Chrome\\Application\\chrome.exe')
    - Open a folder in the file explorer.

    Args:
        path: The full path to the file or application.

    Returns:
        A success or error message.
    """
    try:
        if not os.path.exists(path):
            return f"Error: The path '{path}' does not exist."

        if platform.system() == "Windows":
            os.startfile(path)
        elif platform.system() == "Darwin":  # macOS
            subprocess.run(["open", path], check=True)
        else:  # Linux
            subprocess.run(["xdg-open", path], check=True)
            
        return f"Successfully opened or launched '{path}'."
    except Exception as e:
        return f"Error opening file or application: {str(e)}"


@tool
def Current_time() -> str:
    """Returns the current date and time as a formatted string."""
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

@tool
def GetSystemInfo() -> str:
    """
    Get basic information about the current system of the user laptop or computer.

    Returns:
        str: A string containing OS, version, architecture, processor, and hostname.
    """
    return f"OS: {platform.system()}, \
            OS Version: {platform.version()}, \
            Architecture: {platform.machine()}, \
            Processor: {platform.processor()}, \
            Hostname: {socket.gethostname()}"

@tool
def SeeScreen(ScreenFocus: str = None) -> str:
    """
    Analyzes the screen content using an optimized, high-speed capture method.
    
    This upgraded tool captures the screen, resizes the image for efficiency,
    and uses a vision-language model to provide a detailed description or
    answer a specific question about the on-screen content. It is significantly
    faster and more efficient than a standard screenshot.

    This tool is essential for understanding the current user interface,
    responding to visual prompts, or providing context-aware assistance.

    Args:
        ScreenFocus (str, optional): A specific question or area to focus on 
                                     within the screenshot. If not provided, a 
                                     general description of the screen will be 
                                     returned. Example: "What is the error message?"

    Returns:
        A rich textual description of the screen or a direct answer to the ScreenFocus question.
    """
    try:
        # OPTIMIZATION 1: Use the ultra-fast mss library for screen capture
        with mss.mss() as sct:
            # Get a raw BGRA buffer of the primary monitor
            sct_img = sct.grab(sct.monitors[1])
            # Convert the raw BGRA data to a Pillow Image object
            screenshot = Image.frombytes("RGB", sct_img.size, sct_img.bgra, "raw", "BGRX")

        # OPTIMIZATION 2: Resize the image if it's larger than 1080p
        max_size = (1920, 1080)
        screenshot.thumbnail(max_size, Image.Resampling.LANCZOS)

        # Convert screenshot to base64 data URL (in-memory)
        buffered = BytesIO()
        # OPTIMIZATION 3: Save as JPEG for smaller file size
        screenshot.save(buffered, format="JPEG", quality=90) # Quality set to 90
        img_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
        data_url = f"data:image/jpeg;base64,{img_base64}"

        # Conditionally create the prompt based on whether a focus is provided
        if ScreenFocus:
            system_text = (
                "You are an expert visual assistant. The user has a specific question about the attached screenshot. "
                "Analyze the image and answer their question directly and concisely. Do not describe what you see unless it's the answer. Just provide the answer."
            )
        else:
            system_text = (
                "You are an expert visual assistant. Analyze the following screenshot and provide a clear, structured "
                "description of what is visible. Mention any open applications, UI components, readable text, and "
                "overall screen layout. Be as detailed and organized as possible to help another agent understand "
                "what the user is currently seeing."
            )
        
        system_message = SystemMessage(content=system_text)
        
        human_content = []
        if ScreenFocus:
            human_content.append({"type": "text", "text": ScreenFocus})
        human_content.append({"type": "image_url", "image_url": data_url})
        
        human_message = HumanMessage(content=human_content)

        # Send the structured messages to Gemini and return the description
        response = llm.invoke([system_message, human_message])
        return response.content
    except Exception as e:
        return f"An error occurred while analyzing the screen: {str(e)}"

@tool
def get_active_window_title() -> str:
    """
    Retrieves the title of the currently active (focused) window on the screen.

    This tool is highly efficient for understanding the user's current context, 
    such as which application (e.g., 'Google Chrome', 'File Explorer', 'Visual Studio Code')
    or which specific document ('report-final.docx - Microsoft Word') they are 
    currently working on. It should be used before SeeScreen if the goal is
    just to identify the application the user is in.

    Returns:
        A string containing the title of the active window, or a message if none is found.
    """
    try:
        active_window = gw.getActiveWindow()
        if active_window:
            return f"The current active window is: '{active_window.title}'"
        else:
            # This case is rare on modern OSes but good to have.
            return "No active window found."
    except Exception as e:
        # This can happen if the window closes while the function is running
        return f"An error occurred while getting the active window title: {str(e)}"

@tool
def execute_shell_command(command: str) -> str:
    """
    Executes a shell command in the terminal and returns the output.
    
    This tool is extremely powerful and can be used to run any command-line
    instruction on the user's operating system. It can be used for a wide
    range of tasks, such as:
    - Running scripts (e.g., 'python my_script.py')
    - Managing system processes (e.g., 'tasklist' or 'ps aux')
    - Using developer tools (e.g., 'git status', 'npm install')
    - Getting network information (e.g., 'ipconfig' or 'ifconfig')

    Args:
        command: The command to execute as a string.

    Returns:
        A string containing the standard output and standard error from the command.
    """
    try:
        # For security and to prevent hanging, use a timeout
        result = subprocess.run(
            command,
            shell=True,
            capture_output=True,
            text=True,
            timeout=120,  # 2-minute timeout
            check=False # Do not raise an exception on non-zero exit codes
        )
        output = ""
        if result.stdout:
            output += f"STDOUT:\n{result.stdout}\n"
        if result.stderr:
            output += f"STDERR:\n{result.stderr}\n"
        if not output:
            return "Command executed successfully with no output."
        return output
    except FileNotFoundError:
        return f"Error: Command '{command.split()[0]}' not found. Make sure it is in your system's PATH."
    except subprocess.TimeoutExpired:
        return "Error: Command timed out after 120 seconds."
    except Exception as e:
        return f"An unexpected error occurred: {str(e)}"

@tool
def zip_files(source_path: str, output_zip_path: str) -> str:
    """
    Compresses a single file or an entire directory into a .zip archive.

    Args:
        source_path: The full path to the source file or directory to be zipped.
        output_zip_path: The full path for the output .zip file (e.g., 'C:\\Users\\hp\\Desktop\\archive.zip').

    Returns:
        A success or error message.
    """
    try:
        with zipfile.ZipFile(output_zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            if os.path.isdir(source_path):
                for root, _, files in os.walk(source_path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, start=source_path)
                        zipf.write(file_path, arcname=arcname)
            else: # It's a single file
                zipf.write(source_path, arcname=os.path.basename(source_path))
        return f"Successfully created zip file: {output_zip_path}"
    except Exception as e:
        return f"Error creating zip file: {str(e)}"

@tool
def unzip_file(zip_path: str, destination_dir: str) -> str:
    """
    Extracts the contents of a .zip file to a specified directory.

    Args:
        zip_path: The full path to the .zip file to be extracted.
        destination_dir: The directory where the contents should be extracted.

    Returns:
        A success or error message.
    """
    try:
        os.makedirs(destination_dir, exist_ok=True)
        with zipfile.ZipFile(zip_path, 'r') as zipf:
            zipf.extractall(destination_dir)
        return f"Successfully extracted '{zip_path}' to '{destination_dir}'."
    except Exception as e:
        return f"Error extracting zip file: {str(e)}"

@tool
def type_on_screen(text_to_type: str, interval_seconds: float = 0.05) -> str:
    """
    Types the given text into the currently active window, simulating keyboard input.
    This is useful for filling out forms, writing in text editors, or entering commands.

    Make sure the desired window/input field is in focus before calling this tool.

    Args:
        text_to_type: The string of text to type.
        interval_seconds: The time to wait between each keystroke.

    Returns:
        A success message.
    """
    try:
        pyautogui.write(text_to_type, interval=interval_seconds)
        return "Text typed successfully."
    except Exception as e:
        return f"An error occurred while typing: {str(e)}"

# Create embeddings
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
@tool
def ask_document(content: str, query: str) -> str:
    """
    Analyzes a provided block of text ("the document") to answer a specific query.

    This tool is ideal for situations where you need to answer questions based *only* on a specific
    piece of content, such as the text from a file, a long string from the clipboard, or a previous
    tool's output. It does not use external knowledge.

    Args:
        content: The text content (the "document") to be analyzed and queried.
        query: The specific question to "ask" the document.

    Returns:
        A string containing the answer to the query, derived exclusively from the content.
        Returns an error message if the content is too short or an analysis cannot be performed.
    """

    if not content or not content.strip():
        return "Error: The provided content is empty. Cannot perform analysis."

    # 1. Split the document into chunks
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    docs = text_splitter.create_documents([content])

    # 2. Create an in-memory vector store from the chunks
    vectorstore = Chroma.from_documents(
        documents=docs, 
        embedding=embeddings
    )

    # 3. Create a retriever from the vector store
    retriever = vectorstore.as_retriever(
        search_kwargs={"k": 3},  # Retrieve top 3 chunks
        search_type = "similarity_score_threshold"
    )

    # 4. Define the RAG prompt template
    prompt = PromptTemplate.from_template("""Answer the question based only on the following context:
        {context}

        Question: {question}
        """)
    
    # 5. Create the RAG chain using LangChain Expression Language (LCEL)
    rag_chain = (
        {"context": retriever, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )
    
    # 6. Invoke the chain with the query to get the answer
    answer = rag_chain.invoke(query)

    # Clean up the in-memory vector store
    vectorstore.delete_collection()

    return answer

@tool
def summarize_content(content: str) -> str:
    """
    Creates a concise summary of a given block of text.

    This tool is highly effective for condensing long articles, documents, reports,
    or any other substantial piece of text into its key points. Use it when you
    need to understand the main ideas of a text without reading it in its entirety.

    Args:
        content: The string of text that you want to summarize.

    Returns:
        A string containing a clear and concise summary of the input content.
    """

    if not content or not content.strip():
        return "Error: The provided content is empty and cannot be summarized."

    # Define prompt template
    prompt = PromptTemplate.from_template("""
    You are an expert summarization engine. Your task is to provide a clear and concise summary of the following content.

    Focus on extracting the main ideas, key arguments, and any important conclusions.
    Present the summary in a way that is easy to read and understand.

    Content to summarize:
    ---
    {content}
    ---

    Concise Summary:
    """)
    
    # Create the summarization chain using LangChain Expression Language (LCEL)
    summarizer_chain = (
        {"content": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )
    
    # Invoke the chain with the content to get the summary
    summary = summarizer_chain.invoke(content)

    return summary

@tool
def SmartWebScraper(link: str) -> str:
    """
    Extract meaningful text content from a given webpage,
    focusing on important HTML elements such as paragraphs, headings, lists,
    quotes, and code blocks.

    This tool is optimized for extracting the core readable content from articles,
    documentation, blog posts, and other content-heavy pages.

    Args:
        link (str): The URL of the web page to scrape.

    Returns:
        str: A cleaned and readable concatenation of the page's core textual elements.
    """
    
    try:
        # Add a header to mimic a real browser. This is crucial for avoiding blocks.
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }

        # Automatically detect and use system-configured proxies
        proxies = requests.utils.getproxies()
        response = requests.get(link, headers=headers, proxies=proxies, timeout=10)
        
        response.raise_for_status()  # Raise an exception for bad status codes (e.g., 404 Not Found)
        soup = BeautifulSoup(response.content, "html.parser")

        tags_to_find = ["p", "h1", "h2", "h3", "h4", "h5", "h6", "code"]
        tags = soup.find_all(tags_to_find)
        
        cleaned = []
        for tag in tags:
            text = tag.get_text(strip=True)
            if text:
                cleaned.append(text)

        return "\n\n".join(cleaned)
        
    except Exception as e:
        return f"Error scraping {link}: {str(e)}"

# Create the Tavily search tool
search_tool = TavilySearchResults(max_results=5)



# Combine all tools
all_tools = [get_username, GetSystemInfo, 
        Save_to_clipboard, get_clipboard_content, 
        list_directory_tree, find_files, 
        Delete, Create, Move, Rename,
        Write, Read, zip_files, unzip_file, open_file_or_app,
        search_tool,
        Current_time, SeeScreen, get_active_window_title, 
        execute_shell_command, 
        ask_document, summarize_content,
        SmartWebScraper]


# Create agent 🤖
agent = llm.bind_tools(all_tools)
